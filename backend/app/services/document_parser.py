import io

from docx import Document


def extract_docx_text(file_bytes: bytes) -> tuple[str, list[list[str]]]:
    doc = Document(io.BytesIO(file_bytes))

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)

    full_text = "\n".join(paragraphs)
    return full_text, tables
